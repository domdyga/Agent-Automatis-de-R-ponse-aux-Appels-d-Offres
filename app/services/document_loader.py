from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any

from langchain_core.documents import Document
from app.utils.logger import logger


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _base_metadata(path: Path, extra: dict[str, Any] | None = None) -> dict[str, Any]:
    meta: dict[str, Any] = {
        "source": str(path),
        "filename": path.name,
        "file_type": path.suffix.lstrip(".").lower(),
    }
    if extra:
        meta.update(extra)
    return meta


def load_pdf(path: Path, metadata: dict[str, Any] | None = None) -> list[Document]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf manquant — lance : pip install pypdf") from exc

    reader = PdfReader(str(path))
    docs: list[Document] = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = _clean_text(page.extract_text() or "")
        if not text:
            logger.warning("Page %d vide dans '%s', ignorée", page_num, path.name)
            continue
        meta = _base_metadata(path, metadata)
        meta["page"] = page_num
        docs.append(Document(page_content=text, metadata=meta))

    logger.info("PDF '%s' : %d page(s) chargée(s)", path.name, len(docs))
    return docs


def load_docx(path: Path, metadata: dict[str, Any] | None = None) -> list[Document]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx manquant — lance : pip install python-docx") from exc

    doc = docx.Document(str(path))
    text = _clean_text("\n\n".join(p.text for p in doc.paragraphs if p.text.strip()))

    if not text:
        raise ValueError(f"'{path.name}' ne contient aucun texte.")

    logger.info("DOCX '%s' : %d caractères", path.name, len(text))
    return [Document(page_content=text, metadata=_base_metadata(path, metadata))]


def load_txt(path: Path, metadata: dict[str, Any] | None = None) -> list[Document]:
    text = _clean_text(path.read_text(encoding="utf-8", errors="replace"))
    if not text:
        raise ValueError(f"'{path.name}' est vide.")
    logger.info("TXT '%s' : %d caractères", path.name, len(text))
    return [Document(page_content=text, metadata=_base_metadata(path, metadata))]


def load_csv(path: Path, metadata: dict[str, Any] | None = None) -> list[Document]:
    docs: list[Document] = []
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        for row_num, row in enumerate(csv.DictReader(f), start=1):
            text = _clean_text(" | ".join(f"{k}: {v}" for k, v in row.items() if v))
            if not text:
                continue
            meta = _base_metadata(path, metadata)
            meta["row"] = row_num
            docs.append(Document(page_content=text, metadata=meta))

    if not docs:
        raise ValueError(f"'{path.name}' ne contient aucune donnée.")

    logger.info("CSV '%s' : %d ligne(s)", path.name, len(docs))
    return docs


_LOADERS = {"pdf": load_pdf, "docx": load_docx, "txt": load_txt, "csv": load_csv}
SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


def load_document(path: str | Path, metadata: dict[str, Any] | None = None) -> list[Document]:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {path}")

    ext = path.suffix.lstrip(".").lower()
    loader = _LOADERS.get(ext)
    if not loader:
        raise ValueError(f"Format '.{ext}' non supporté. Acceptés : {sorted(SUPPORTED_EXTENSIONS)}")

    return loader(path, metadata)
